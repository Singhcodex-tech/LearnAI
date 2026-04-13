"""
LAW ASSIGNMENT GENERATOR — Flask Blueprint
==========================================
Isolated feature. Uses host app's _call_groq, GROQ_API_KEY, limiter.
Mount in app.py:
    from features.assignment.blueprint import assignment_bp
    app.register_blueprint(assignment_bp)

Routes added:
    POST /api/assignment/generate
    GET  /api/assignment/<id>/docx
    GET  /api/assignment/<id>/pdf
    GET  /api/assignment/<id>/status

Deps (add to requirements.txt):
    python-docx
    reportlab
"""

from __future__ import annotations
import io
import re
import time
import uuid
from datetime import datetime
from typing import Optional

from flask import Blueprint, jsonify, request, send_file

# ── Blueprint ────────────────────────────────────────────────────────────────
assignment_bp = Blueprint("assignment", __name__, url_prefix="/api/assignment")

# ── In-module store (no global pollution) ───────────────────────────────────
_STORE: dict = {}          # { id: { meta, sections, created_at } }
_STORE_TTL   = 7200        # 2 hours

def _prune():
    now = time.time()
    stale = [k for k, v in _STORE.items() if now - v["created_at"] > _STORE_TTL]
    for k in stale:
        del _STORE[k]

def _find_cached(topic: str) -> Optional[dict]:
    for v in _STORE.values():
        if v["meta"]["topic"].lower().strip() == topic.lower().strip():
            return v
    return None

# ── Token budgets per section ────────────────────────────────────────────────
_TOKEN = {
    "declaration":     250,
    "acknowledgment":  200,
    "introduction":    500,
    "body":            950,
    "conclusion":      320,
    "bibliography":    420,
}

# ── Prompts ──────────────────────────────────────────────────────────────────
def _prompt(section: str, topic: str, meta: dict) -> str:
    base = f'Write {section} for a law assignment on: "{topic}". Formal, academic, structured. Include legal reasoning. Use [CASE: brief description] placeholders for uncertain case citations.'
    prompts = {
        "declaration": (
            f'Write a 1-paragraph student academic honesty declaration for a law assignment.\n'
            f'Student: {meta["student_name"]}, Roll No: {meta["roll_no"]}, '
            f'Course: {meta["course"]}, College: {meta["college"]}.\n'
            f'Formal tone. Max 80 words.'
        ),
        "acknowledgment": (
            f'Write a brief acknowledgment (80-100 words) for a law assignment by '
            f'{meta["student_name"]} of {meta["college"]}. '
            f'Thank faculty guide, library, and family. Formal academic tone.'
        ),
        "introduction": (
            f'{base}\n'
            f'Cover: definition, constitutional/statutory basis, scope, significance, research objectives. '
            f'250-300 words. End with a clear thesis statement.'
        ),
        "body": (
            f'{base}\n'
            f'Structure with 3-4 numbered sub-sections (e.g. "2.1 Constitutional Framework").\n'
            f'Include: statutory provisions, judicial interpretation, critical analysis.\n'
            f'Insert [CASE: description] where case law is relevant but uncertain.\n'
            f'600-750 words total.'
        ),
        "conclusion": (
            f'{base}\n'
            f'Summarise key findings, legal implications, policy recommendations. '
            f'150-180 words. No new arguments.'
        ),
        "bibliography": (
            f'List 6-8 bibliography entries in Bluebook citation format for a law assignment on: "{topic}".\n'
            f'Mix of: Supreme Court cases, statutes, legal textbooks, journal articles.\n'
            f'Use [CITATION NEEDED] where exact details are uncertain.\n'
            f'Number each entry. No extra text.'
        ),
    }
    return prompts.get(section, base)

# ── Section generator (uses host _call_groq) ─────────────────────────────────
def _gen_section(section: str, topic: str, meta: dict) -> str:
    from app import _call_groq  # lazy import — no circular dep at module load
    system = (
        "You are a senior legal academic writer at a reputed Indian law school. "
        "Write in formal, precise academic English. Use correct legal terminology. "
        "Structure content clearly. Never use markdown symbols like ** or ##."
    )
    result = _call_groq(
        _prompt(section, topic, meta),
        max_tokens=_TOKEN.get(section, 500),
        system=system,
    )
    return (result or "").strip()

# ── Assembly ──────────────────────────────────────────────────────────────────
def _assemble(meta: dict) -> dict:
    _prune()
    cached = _find_cached(meta["topic"])
    if cached:
        return cached

    sections = {}
    order = ["declaration"]
    if meta.get("include_acknowledgment", True):
        order.append("acknowledgment")
    order += ["introduction", "body", "conclusion", "bibliography"]

    for sec in order:
        sections[sec] = _gen_section(sec, meta["topic"], meta)

    record = {
        "id":         str(uuid.uuid4()),
        "meta":       meta,
        "sections":   sections,
        "created_at": time.time(),
    }
    _STORE[record["id"]] = record
    return record

# ── DOCX builder ──────────────────────────────────────────────────────────────
def _build_docx(record: dict) -> bytes:
    try:
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError:
        raise RuntimeError("python-docx not installed. Run: pip install python-docx")

    meta     = record["meta"]
    sections = record["sections"]
    doc      = Document()

    # ── Page margins ────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin   = Inches(1.25)
        section.right_margin  = Inches(1.0)

    # ── Helpers ─────────────────────────────────────────────────────────────
    def set_font(run, size=12, bold=False, color=None):
        run.font.name = "Times New Roman"
        run.font.size = Pt(size)
        run.bold      = bold
        if color:
            run.font.color.rgb = RGBColor(*color)

    def add_heading(text, level=1):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(18 if level == 1 else 12)
        p.paragraph_format.space_after  = Pt(8)
        run = p.add_run(text)
        set_font(run, size=14 if level == 1 else 12, bold=True,
                 color=(26, 58, 92) if level == 1 else (46, 95, 138))
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        return p

    def add_body(text):
        if not text:
            return
        for line in text.split("\n"):
            line = line.strip()
            if not line:
                continue
            # Detect sub-headings: "2.1 ..." or "I. ..."
            is_subhead = bool(re.match(r'^(\d+\.\d+|[IVX]+\.)\s', line))
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(line)
            if is_subhead:
                set_font(run, size=12, bold=True, color=(46, 95, 138))
            else:
                set_font(run, size=12)
            p.paragraph_format.alignment = (
                WD_ALIGN_PARAGRAPH.LEFT if is_subhead else WD_ALIGN_PARAGRAPH.JUSTIFY
            )

    def add_divider():
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after  = Pt(4)
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "2E4057")
        pBdr.append(bottom)
        pPr.append(pBdr)

    def page_break():
        doc.add_page_break()

    # ── COVER PAGE ───────────────────────────────────────────────────────────
    for _ in range(4):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(meta["college"].upper())
    set_font(run, 16, bold=True, color=(26, 58, 92))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(meta["course"])
    set_font(run, 13)

    doc.add_paragraph()
    add_divider()
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("LAW ASSIGNMENT")
    set_font(run, 20, bold=True, color=(26, 58, 92))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("ON")
    set_font(run, 12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(meta["topic"].upper())
    set_font(run, 15, bold=True, color=(26, 58, 92))

    doc.add_paragraph()
    add_divider()
    for _ in range(3):
        doc.add_paragraph()

    for label, value in [
        ("Submitted by", meta["student_name"]),
        ("Roll No",      meta["roll_no"]),
        ("Date",         datetime.now().strftime("%d %B %Y")),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = p.add_run(f"{label}: ")
        set_font(r1, 12)
        r2 = p.add_run(value)
        set_font(r2, 12, bold=True)

    page_break()

    # ── TABLE OF CONTENTS (manual) ───────────────────────────────────────────
    add_heading("TABLE OF CONTENTS")
    add_divider()
    toc_entries = [
        ("Declaration",                            "2"),
        ("Acknowledgment",                         "3") if sections.get("acknowledgment") else None,
        ("1.  Introduction",                       "4"),
        ("2.  Legal Analysis",                     "5"),
        ("3.  Conclusion",                         "8"),
        ("4.  Bibliography",                       "9"),
    ]
    for entry in toc_entries:
        if not entry:
            continue
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        r1 = p.add_run(entry[0])
        set_font(r1, 12)
        r2 = p.add_run(f"{'.' * max(1, 55 - len(entry[0]))} {entry[1]}")
        set_font(r2, 12)
    page_break()

    # ── DECLARATION ──────────────────────────────────────────────────────────
    add_heading("DECLARATION")
    add_divider()
    add_body(sections.get("declaration", ""))
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run(f"Signature: ________________        Date: {datetime.now().strftime('%d %B %Y')}")
    set_font(run, 11)
    page_break()

    # ── ACKNOWLEDGMENT ───────────────────────────────────────────────────────
    if sections.get("acknowledgment"):
        add_heading("ACKNOWLEDGMENT")
        add_divider()
        add_body(sections["acknowledgment"])
        page_break()

    # ── INTRODUCTION ─────────────────────────────────────────────────────────
    add_heading("1.  INTRODUCTION")
    add_divider()
    add_body(sections.get("introduction", ""))
    page_break()

    # ── BODY ─────────────────────────────────────────────────────────────────
    add_heading("2.  LEGAL ANALYSIS")
    add_divider()
    add_body(sections.get("body", ""))
    page_break()

    # ── CONCLUSION ───────────────────────────────────────────────────────────
    add_heading("3.  CONCLUSION")
    add_divider()
    add_body(sections.get("conclusion", ""))
    page_break()

    # ── BIBLIOGRAPHY ─────────────────────────────────────────────────────────
    add_heading("4.  BIBLIOGRAPHY")
    p = doc.add_paragraph()
    run = p.add_run("(Bluebook Citation Format)")
    set_font(run, 10, color=(100, 100, 100))
    add_divider()
    add_body(sections.get("bibliography", ""))

    # ── Save to bytes ────────────────────────────────────────────────────────
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ── PDF builder ───────────────────────────────────────────────────────────────
def _build_pdf(record: dict) -> bytes:
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.lib import colors
        from reportlab.platypus import (
            SimpleDocTemplate, Paragraph, Spacer, PageBreak, HRFlowable
        )
        from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT
    except ImportError:
        raise RuntimeError("reportlab not installed. Run: pip install reportlab")

    meta     = record["meta"]
    sections = record["sections"]
    buf      = io.BytesIO()

    doc = SimpleDocTemplate(
        buf, pagesize=A4,
        leftMargin=1.25*inch, rightMargin=inch,
        topMargin=inch, bottomMargin=inch,
    )

    NAVY  = colors.HexColor("#1A3A5C")
    BLUE  = colors.HexColor("#2E5F8A")
    GREY  = colors.HexColor("#64748B")
    BLACK = colors.black

    styles = getSampleStyleSheet()
    S = {
        "cover_college": ParagraphStyle("cc", fontName="Helvetica-Bold", fontSize=15,
                                         textColor=NAVY, alignment=TA_CENTER, spaceAfter=8),
        "cover_course":  ParagraphStyle("cs", fontName="Helvetica", fontSize=12,
                                         textColor=BLACK, alignment=TA_CENTER, spaceAfter=20),
        "cover_title":   ParagraphStyle("ct", fontName="Helvetica-Bold", fontSize=18,
                                         textColor=NAVY, alignment=TA_CENTER, spaceAfter=10),
        "cover_sub":     ParagraphStyle("csb", fontName="Helvetica", fontSize=11,
                                         textColor=BLACK, alignment=TA_CENTER, spaceAfter=6),
        "cover_name":    ParagraphStyle("cn", fontName="Helvetica-Bold", fontSize=12,
                                         textColor=BLACK, alignment=TA_CENTER, spaceAfter=6),
        "h1":  ParagraphStyle("h1", fontName="Helvetica-Bold", fontSize=14,
                               textColor=NAVY, spaceBefore=18, spaceAfter=8),
        "h2":  ParagraphStyle("h2", fontName="Helvetica-Bold", fontSize=12,
                               textColor=BLUE, spaceBefore=12, spaceAfter=6),
        "body": ParagraphStyle("body", fontName="Times-Roman", fontSize=11,
                                textColor=BLACK, alignment=TA_JUSTIFY,
                                spaceBefore=0, spaceAfter=6, leading=16),
        "toc":  ParagraphStyle("toc", fontName="Helvetica", fontSize=11,
                                textColor=BLACK, spaceAfter=4),
        "sig":  ParagraphStyle("sig", fontName="Helvetica", fontSize=10,
                                textColor=GREY, spaceAfter=4),
        "biblio": ParagraphStyle("bib", fontName="Times-Roman", fontSize=10,
                                  textColor=BLACK, spaceAfter=5, leading=14),
    }

    def hr(): return HRFlowable(width="100%", thickness=1, color=NAVY, spaceAfter=8, spaceBefore=4)

    story = []

    # ── Cover ────────────────────────────────────────────────────────────────
    story += [Spacer(1, 1.2*inch)]
    story.append(Paragraph(meta["college"].upper(), S["cover_college"]))
    story.append(Paragraph(meta["course"], S["cover_course"]))
    story.append(hr())
    story += [Spacer(1, 0.2*inch)]
    story.append(Paragraph("LAW ASSIGNMENT", S["cover_title"]))
    story.append(Paragraph("ON", S["cover_sub"]))
    story.append(Paragraph(f'<b>{meta["topic"].upper()}</b>', S["cover_title"]))
    story += [Spacer(1, 0.3*inch)]
    story.append(hr())
    story += [Spacer(1, 0.8*inch)]
    story.append(Paragraph("Submitted by", S["cover_sub"]))
    story.append(Paragraph(meta["student_name"], S["cover_name"]))
    story.append(Paragraph(f'Roll No: {meta["roll_no"]}', S["cover_sub"]))
    story.append(Paragraph(datetime.now().strftime("%d %B %Y"), S["cover_sub"]))
    story.append(PageBreak())

    # ── TOC ──────────────────────────────────────────────────────────────────
    story.append(Paragraph("TABLE OF CONTENTS", S["h1"]))
    story.append(hr())
    toc_items = [
        ("Declaration", "2"),
        ("Acknowledgment", "3") if sections.get("acknowledgment") else None,
        ("1.  Introduction", "4"),
        ("2.  Legal Analysis", "5"),
        ("3.  Conclusion", "8"),
        ("4.  Bibliography", "9"),
    ]
    for item in toc_items:
        if item:
            dots = "." * max(1, 60 - len(item[0]))
            story.append(Paragraph(f'{item[0]} <font color="#64748B">{dots}</font> {item[1]}', S["toc"]))
    story.append(PageBreak())

    def add_section(title, text, style_key="body"):
        story.append(Paragraph(title, S["h1"]))
        story.append(hr())
        if not text:
            return
        for line in text.split("\n"):
            line = line.strip()
            if not line:
                story.append(Spacer(1, 4))
                continue
            is_sub = bool(re.match(r'^(\d+\.\d+|[IVX]+\.)\s', line))
            sty = S["h2"] if is_sub else S[style_key]
            # Escape < > & for reportlab
            safe = line.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            story.append(Paragraph(safe, sty))

    # ── Declaration ──────────────────────────────────────────────────────────
    add_section("DECLARATION", sections.get("declaration", ""))
    story.append(Spacer(1, 0.2*inch))
    story.append(Paragraph(
        f'Signature: ________________&nbsp;&nbsp;&nbsp;&nbsp;Date: {datetime.now().strftime("%d %B %Y")}',
        S["sig"]
    ))
    story.append(PageBreak())

    # ── Acknowledgment ───────────────────────────────────────────────────────
    if sections.get("acknowledgment"):
        add_section("ACKNOWLEDGMENT", sections["acknowledgment"])
        story.append(PageBreak())

    # ── Sections ─────────────────────────────────────────────────────────────
    add_section("1.  INTRODUCTION",   sections.get("introduction", ""))
    story.append(PageBreak())
    add_section("2.  LEGAL ANALYSIS", sections.get("body", ""))
    story.append(PageBreak())
    add_section("3.  CONCLUSION",     sections.get("conclusion", ""))
    story.append(PageBreak())

    # Bibliography
    story.append(Paragraph("4.  BIBLIOGRAPHY", S["h1"]))
    story.append(Paragraph("(Bluebook Citation Format)", S["sig"]))
    story.append(hr())
    for line in (sections.get("bibliography") or "").split("\n"):
        line = line.strip()
        if line:
            safe = line.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            story.append(Paragraph(safe, S["biblio"]))

    doc.build(story)
    buf.seek(0)
    return buf.read()


# ── Routes ────────────────────────────────────────────────────────────────────

@assignment_bp.route("/generate", methods=["POST"])
def generate_assignment():
    data = request.get_json(silent=True) or {}
    required = ["topic", "student_name", "roll_no", "course", "college"]
    missing  = [f for f in required if not str(data.get(f, "")).strip()]
    if missing:
        return jsonify({"error": f"Missing fields: {', '.join(missing)}"}), 400

    topic = str(data["topic"]).strip()
    if len(topic) > 300:
        return jsonify({"error": "Topic too long (max 300 chars)"}), 400

    meta = {
        "topic":                topic,
        "student_name":         str(data["student_name"]).strip(),
        "roll_no":              str(data["roll_no"]).strip(),
        "course":               str(data["course"]).strip(),
        "college":              str(data["college"]).strip(),
        "include_acknowledgment": bool(data.get("include_acknowledgment", True)),
    }

    try:
        record = _assemble(meta)
    except Exception as e:
        return jsonify({"error": str(e)}), 500

    return jsonify({
        "success":       True,
        "assignment_id": record["id"],
        "sections":      list(record["sections"].keys()),
        "cached":        record.get("created_at", 0) < time.time() - 2,
    })


@assignment_bp.route("/<aid>/docx", methods=["GET"])
def download_docx(aid):
    record = _STORE.get(aid)
    if not record:
        return jsonify({"error": "Assignment not found or expired"}), 404
    try:
        buf = _build_docx(record)
    except RuntimeError as e:
        return jsonify({"error": str(e)}), 500
    safe = re.sub(r'[^a-zA-Z0-9_-]', '_', record["meta"]["topic"][:40]).strip('_')
    return send_file(
        io.BytesIO(buf),
        as_attachment=True,
        download_name=f"{safe}_Assignment.docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


@assignment_bp.route("/<aid>/pdf", methods=["GET"])
def download_pdf(aid):
    record = _STORE.get(aid)
    if not record:
        return jsonify({"error": "Assignment not found or expired"}), 404
    try:
        buf = _build_pdf(record)
    except RuntimeError as e:
        return jsonify({"error": str(e)}), 500
    safe = re.sub(r'[^a-zA-Z0-9_-]', '_', record["meta"]["topic"][:40]).strip('_')
    return send_file(
        io.BytesIO(buf),
        as_attachment=True,
        download_name=f"{safe}_Assignment.pdf",
        mimetype="application/pdf",
    )


@assignment_bp.route("/<aid>/status", methods=["GET"])
def assignment_status(aid):
    record = _STORE.get(aid)
    if not record:
        return jsonify({"error": "Not found"}), 404
    return jsonify({
        "id":         record["id"],
        "topic":      record["meta"]["topic"],
        "student":    record["meta"]["student_name"],
        "sections":   list(record["sections"].keys()),
        "created_at": record["created_at"],
    })
